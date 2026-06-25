import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.config import Settings
from app.services.rag.types import ExtractedBlock, ExtractionResult


class DocumentExtractor:
    def __init__(self, settings: Settings):
        self.settings = settings

    def extract(self, file_path: Path, filename: str, mime_type: str) -> ExtractionResult:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(file_path, mime_type)
        if suffix == ".docx":
            return self._extract_docx(file_path, mime_type)
        raise ValueError("Only PDF and DOCX uploads are supported.")

    def _extract_pdf(self, file_path: Path, mime_type: str) -> ExtractionResult:
        reader = PdfReader(str(file_path))
        blocks: list[ExtractedBlock] = []
        ocr_used = False
        ocr_unavailable = False

        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            used_for_page = False

            if self._should_ocr(text):
                ocr_text, unavailable = self._ocr_pdf_page(file_path, index - 1)
                ocr_unavailable = ocr_unavailable or unavailable
                if ocr_text.strip():
                    text = ocr_text
                    used_for_page = True
                    ocr_used = True

            blocks.append(
                ExtractedBlock(
                    text=self._clean_text(text),
                    page_number=index,
                    ocr_used=used_for_page,
                )
            )

        return ExtractionResult(
            blocks=blocks,
            page_count=len(reader.pages),
            mime_type=mime_type or "application/pdf",
            ocr_enabled=self.settings.rag_enable_ocr,
            ocr_used=ocr_used,
            ocr_unavailable=ocr_unavailable,
        )

    def _extract_docx(self, file_path: Path, mime_type: str) -> ExtractionResult:
        document = Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return ExtractionResult(
            blocks=[ExtractedBlock(text=self._clean_text("\n".join(parts)), page_number=None)],
            page_count=1,
            mime_type=mime_type
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ocr_enabled=False,
            ocr_used=False,
            ocr_unavailable=False,
        )

    def _should_ocr(self, text: str) -> bool:
        return self.settings.rag_enable_ocr and len(text.strip()) < self.settings.rag_ocr_min_chars

    def _ocr_pdf_page(self, file_path: Path, page_index: int) -> tuple[str, bool]:
        try:
            import pypdfium2 as pdfium
            import pytesseract

            if self.settings.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = self.settings.tesseract_cmd

            pdf = pdfium.PdfDocument(str(file_path))
            page = pdf[page_index]
            image = page.render(scale=2).to_pil()
            return pytesseract.image_to_string(image), False
        except Exception:
            return "", True

    def _clean_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
